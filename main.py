# -*- coding: utf-8 -*-
# ======================================================
# MGA IA WEB – FORMULADOR PROFESIONAL
# Fundación Almagua
# ======================================================

import os, json, re, io, zipfile, csv
import pandas as pd

from fastapi import FastAPI, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

from docx import Document

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ======================================================
# 🌐 APP
# ======================================================
app = FastAPI(title="MGA IA – Fundación Almagua", version="1.0")
app.mount("/assets", StaticFiles(directory="assets"), name="assets")

# ======================================================
# 🔐 CONFIGURACIÓN
# ======================================================
MODEL = "gpt-4o-mini"
TEMPERATURE = 0.1

BASE = "data"
PLAN_CSV = f"{BASE}/plan1.csv"

TEXTOS = f"{BASE}/textos"
GUIA_TXT = f"{TEXTOS}/guia.txt"

PDF_BASE = f"{BASE}/pdf_referencia"

llm = ChatOpenAI(model=MODEL, temperature=TEMPERATURE)

# ======================================================
# 📚 RAG PDFs
# ======================================================
def cargar_corpus():
    documentos = []
    if os.path.exists(PDF_BASE):
        for archivo in os.listdir(PDF_BASE):
            if archivo.lower().endswith(".pdf"):
                documentos.extend(PyPDFLoader(os.path.join(PDF_BASE, archivo)).load())
    if not documentos:
        return None
    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=120)
    chunks = splitter.split_documents(documentos)
    return FAISS.from_documents(chunks, OpenAIEmbeddings())

db = cargar_corpus()

# ======================================================
# 🧹 UTILIDADES
# ======================================================
def limpiar_codigo(codigo: str) -> str:
    return re.sub(r"[^0-9]", "", codigo).strip()

def cargar_txt(path: str, max_chars=2500) -> str:
    if not os.path.exists(path):
        return ""
    with open(path, encoding="utf-8") as f:
        return f.read()[:max_chars]

def normalizar_columnas(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = (
        df.columns
        .str.strip()
        .str.lower()
        .str.replace(r"[^\w\s]", "", regex=True)
        .str.replace("á","a").str.replace("é","e")
        .str.replace("í","i").str.replace("ó","o").str.replace("ú","u")
    )
    return df

# ======================================================
# 📊 CONSULTA CSV POR CÓDIGO
# ======================================================
def consultar_plan_por_producto(codigo_producto: str) -> dict:
    if not os.path.exists(PLAN_CSV):
        raise FileNotFoundError("No existe plan1.csv")
    
    df_raw = pd.read_csv(PLAN_CSV, header=None, encoding="utf-8-sig")

    header_row = None
    for i, row in df_raw.iterrows():
        fila_texto = " ".join(str(c).lower().replace("á","a").replace("é","e")
                               .replace("í","i").replace("ó","o").replace("ú","u")
                               .strip() for c in row if pd.notna(c))
        if "codigo" in fila_texto and "producto" in fila_texto:
            header_row = i
            break

    if header_row is None:
        raise ValueError("No se pudo detectar la fila de encabezados en el CSV")

    df = pd.read_csv(PLAN_CSV, header=header_row, encoding="utf-8-sig")
    df = normalizar_columnas(df)

    col_codigo = next((c for c in df.columns if "codigo" in c and "producto" in c), None)
    if col_codigo is None:
        raise ValueError(f"No se encontró columna de código. Columnas: {df.columns.tolist()}")

    df[col_codigo] = df[col_codigo].astype(str).str.strip()
    fila = df[df[col_codigo] == codigo_producto]
    if fila.empty:
        return {}

    r = fila.iloc[0]

    # Campos limpios para encabezado del documento
    return {
        "codigo_producto": codigo_producto,
        "sector": r.get("sector",""),
        "programa": r.get("programa presupuestal",""),
        "producto": r.get("producto",""),
        "descripcion_producto": r.get("descripcion del producto",""),
        "objetivo_linea": r.get("objetivo linea estrategica",""),
        "indicador_producto": r.get("indicador de producto",""),
        "meta_producto": r.get("meta de producto programada al cuatrienio",""),
        "secretaria": r.get("secretaria  dependencia responsable del  producto",""),
        "formulador": "Iván Darío Pacheco Sierra",  # campo fijo limpio
        "fecha_creacion": "15/01/2026"               # fecha fija o actualizable
    }

# ======================================================
# 🧠 REDACCIÓN DOCUMENTO TÉCNICO MGA COMPLETO
# ======================================================
def redactar_documento_tecnico_completo(datos_plan: dict) -> str:
    """
    Genera un documento técnico MGA completo en formato texto,
    siguiendo la guía metodológica y ejemplo, con más argumentos,
    estructura formal y coherente.
    """
    contexto_pdf = ""
    if db and datos_plan.get("producto"):
        # buscar PDFs relacionados como referencia (máx 3 documentos)
        docs = db.similarity_search(datos_plan["producto"], k=3)
        contexto_pdf = "\n".join(d.page_content for d in docs)[:4000]  # más caracteres

    prompt = f"""
Eres un FORMULADOR PROFESIONAL MGA – COLOMBIA. 
Redacta un DOCUMENTO TÉCNICO MGA completo para el proyecto con código {datos_plan.get('codigo_producto')}.

### ENCABEZADO DEL PROYECTO
Nombre del proyecto: {datos_plan.get('producto')}
Código BPIN: {datos_plan.get('codigo_producto')}
Sector: {datos_plan.get('sector')}
Programa: {datos_plan.get('programa')}
Fecha de creación: {datos_plan.get('fecha_creacion')}
Formulador: {datos_plan.get('formulador')}

### OBJETIVO DEL DOCUMENTO
Redacta un documento técnico detallado que incluya:
- Contexto del proyecto
- Diagnóstico y análisis del problema
- Árbol de problemas y causas principales
- Cadena de valor y actores clave
- Estrategias y soluciones propuestas
- Indicadores de producto y resultados
- Metodología y plan de implementación
- Referencias técnicas (sin inventar datos)

### LINEAMIENTOS DE REDACCIÓN
- Usa la guía MGA (guia.txt) como referencia para estructura y formato
- Mantén coherencia, formalidad y claridad
- Incluye todos los elementos metodológicos: árbol de problemas, cadena de valor, plan de acción, indicadores
- Extiende la redacción con explicaciones y argumentos sólidos
- No inventes información que no esté soportada por el proyecto o la guía

### REFERENCIA TÉCNICA
{contexto_pdf}

Redacta el DOCUMENTO TÉCNICO MGA completo, listo para ser utilizado como documento oficial.
"""

    # Invocar LLM con mayor contexto y longitud de salida
    return llm.invoke(prompt).content


# ======================================================
# 📄 GENERAR ARCHIVOS
# ======================================================
def generar_docx(texto: str) -> bytes:
    doc = Document()
    doc.add_heading("DOCUMENTO TÉCNICO MGA", 0)
    for p in texto.split("\n\n"):
        doc.add_paragraph(p)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generar_txt(datos_plan: dict) -> bytes:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Código Producto", "Producto", "Descripción", "Sector", "Programa", "Formulador", "Fecha"])
    writer.writerow([
        datos_plan.get("codigo_producto",""),
        datos_plan.get("producto",""),
        datos_plan.get("descripcion_producto",""),
        datos_plan.get("sector",""),
        datos_plan.get("programa",""),
        datos_plan.get("formulador",""),
        datos_plan.get("fecha_creacion","")
    ])
    return buffer.getvalue().encode("utf-8")

def generar_zip(codigo: str, docx: bytes, txt: bytes) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"MGA_{codigo}.docx", docx)
        z.writestr(f"MGA_{codigo}.csv", txt)
    buffer.seek(0)
    return buffer.read()

# ======================================================
# 🌐 WEB
# ======================================================
@app.get("/", response_class=HTMLResponse)
def home():
    return open("index.html", encoding="utf-8").read()

@app.post("/generar")
def generar(codigo_producto: str = Form(...)):
    codigo_producto = limpiar_codigo(codigo_producto)
    datos = consultar_plan_por_producto(codigo_producto)

    if not datos:
        return HTMLResponse(f"<h2>No existe información para el código {codigo_producto}</h2>")

    # ✅ Llamada corregida a la función completa
    texto = redactar_documento_tecnico_completo(datos)
    docx = generar_docx(texto)
    txt = generar_txt(datos)
    zip_bytes = generar_zip(codigo_producto, docx, txt)

    return StreamingResponse(
        io.BytesIO(zip_bytes),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=MGA_{codigo_producto}.zip"}
    )
